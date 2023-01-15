import dash
import dash
from dash import html
from dash import dcc
import dash_bootstrap_components as dbc
from dash import Input, Output, State, ctx
import base64
import docx
import io
from resume_parse import extract_experiences, extract_projects, extract_skills
from completions import rewrite_resume, get_keywords, keywords_to_list
from embeddings import get_embeddings, get_cosine_similarity
import math

x = ''
app = dash.Dash(__name__, external_stylesheets=[dbc.themes.UNITED])
app.layout = dbc.Container(id = "view1",children=[
    html.H1('Resume Optimiser', style={'textAlign':'center'}),
    html.H4('Powered by GPT-3', style={'textAlign':'center'}),
    html.Br(),
    html.Br(),
    dbc.Container(
        [
            dbc.Row(
                [
                    dbc.Col(id = "col1", children = [html.Div(children = [
                        dcc.Upload(id = "file_upload", children = html.Div(['Drag and Drop or ',
            html.A('Select Files')])),
                        html.Div(id="file-name",children="No File Uploaded"),
                        html.Br(),
                     dbc.Textarea(id="job-description", className="text-area", placeholder="Copy Job Description Here")
                    ])]),
                    dbc.Col(id = "boost", children=[dbc.Label("Boost Meter", html_for="slider"),
                    dcc.Slider(id="boost_score", min=0, max=10, step=1, value=3),
                    html.Div([dbc.Button("Boost Resume",color="primary",id="boost-btn",n_clicks=0),
                    html.Br(),
                    html.Br(),

                    dbc.Spinner(children=[html.Div(id="loading")],color="primary")])
                    
                    ]),
                    # dbc.Col(html.Div(children = [dbc.Select(id="filter", value='1', options =[{"label": "Experience", "value": "1"}, {"label": "Skills", "value": "2"},{"label": "Projects", "value": "3"}]), 
                    dbc.Col(html.Div(children = [dbc.Tabs(id="filter", active_tab='tab-1', children =[
                        dbc.Tab(label="Experience", tab_id="tab-1"),
                        dbc.Tab(label="Skills", tab_id="tab-2"),
                        dbc.Tab(label="Projects", tab_id="tab-3")
                        #{"label": "Experience", "value": "1"}, {"label": "Skills", "value": "2"},{"label": "Projects", "value": "3"}
                        ]), 
                    dbc.Textarea(id="old_area", className='text-area mt-2', placeholder="Before"),dbc.Textarea(id="new_area", className='text-area mt-2', placeholder="✨ After ✨"), 
                    html.Div([dbc.Button('Download',className='mt-2', id="dl-btn",n_clicks=0), dcc.Download(id="download-text")])]))
                ]
            )
        ]
    ,fluid=True) 
], fluid=True)
def parse_contents(contents, filename, jd, boost_score):
    content_type, content_string = contents.split(',')
    decoded = base64.b64decode(content_string)
    if 'docx' in filename:
        # read docx
        doc = docx.Document(io.BytesIO(decoded))
        parsed_resume = '\n\n'.join([paragraph.text for paragraph in doc.paragraphs])
        # extract experiences, skills and projects 
        experience = extract_experiences(parsed_resume)
        skills = extract_skills(parsed_resume)
        projects = extract_projects(parsed_resume)
        resume_dict = {}
        resume_dict['experience'] = experience
        resume_dict['skills'] = skills
        resume_dict['projects'] = projects
        # use gpt-3 to extract keywords from jd 
        keywords = keywords_to_list(get_keywords(jd))
        # sort keywords based on cosine similarity to the combination of experience, skills, and projects 
        combined_embeddings = get_embeddings(experience + skills + projects)
        sorted_keywords = sorted(keywords, key=lambda x: get_cosine_similarity(combined_embeddings, get_embeddings(x)), reverse=True)
        # depending on boost_score, adjust number of keywords to use starting from keyword with highest cosine similarity
        final_keywords = sorted_keywords[ : int(boost_score / 10 * len(sorted_keywords))]
        # print("ORIGINAL KEYWORDS", keywords)
        # print("SORTED KEYWORDS", sorted_keywords)
        # print("FINAL KEYWORDS", final_keywords)
        # use keywords to rewrite resume
        result = rewrite_resume(resume_dict, ", ".join(final_keywords), boost_score)
        #print(result)
        return experience, skills, projects, result
        #return  'NEW EXPERIENCE: \n' + result.get('experience')  + '\n\n NEW SKILLS: \n' + result.get('skills') +'\n\n NEW PROJECTS: \n' + result.get('projects')
    else:
        return 'Invalid file type'

@app.callback(
    Output("file-name", 'children'),
    Input('file_upload', 'filename')
)

def show_upload_name(filename):
    if filename:
        return filename
    else:
        return "No File Uploaded"

@app.callback(
    Output('old_area','value'),
    Output('new_area','value'),
    Output('loading','children'),
    Input('boost-btn','n_clicks'),
    Input("filter",'active_tab'),
    State('file_upload', 'contents'),
    State('file_upload', 'filename'),
    State('job-description','value'),
    State('boost_score','value')
    )

def update_output(n_clicks, filter, contents, filename, jd, boost_score):
    triggered_id  = ctx.triggered_id
    global x
    if triggered_id == 'boost-btn':
        x = parse_contents(contents, filename, jd, boost_score)
        #old_output = f'OLD EXPERIENCE: {x[0]}'+ f'\n\n OLD SKILLS: {x[1]}' + f'\n\n OLD PROJECTS: {x[2]}'
        #new_output = 'NEW EXPERIENCE: \n' + x[3].get('experience')  + '\n\n NEW SKILLS: \n' + x[3].get('skills') +'\n\n NEW PROJECTS: \n' + x[3].get('projects')
        return x[0],x[3].get('experience'),"Done!"
    elif triggered_id == 'filter':
        if filter == 'tab-1':
            return x[0],x[3].get('experience'),"Done!"
        if filter =='tab-2':
            return x[1],x[3].get('skills'),"Done!"
        if filter == 'tab-3':
            return x[2],x[3].get('projects'),"Done!"

@app.callback(
    Output('download-text','data'),
    Input('dl-btn','n_clicks'),
    State('new_area','value')
)

def download_file(n_clicks,value):
    if n_clicks:
        mydoc = docx.Document()
        mydoc.add_paragraph(value)
        mydoc.save("new_resume.docx")
        return dcc.send_file('new_resume.docx')

if __name__ == "__main__":
    app.run_server(debug=True)
